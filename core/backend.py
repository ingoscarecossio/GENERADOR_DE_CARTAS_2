
# -*- coding: utf-8 -*-
from __future__ import annotations
import io, re
from datetime import datetime
from typing import Dict, Optional, List
import pandas as pd
from docx import Document
from unidecode import unidecode

def _norm(s: str) -> str:
    s = unidecode(str(s or "").strip().lower())
    s = re.sub(r"\s+", " ", s)
    return s

def slugify(text: str) -> str:
    s = unidecode(str(text or "").strip())
    s = re.sub(r"[^A-Za-z0-9_\- ]+", "", s)
    s = s.strip().replace(" ", "_")
    return s or "SIN_GRUPO"

SYNONYMS = {
    "actor": ["actor", "columna a", "a", "interesado", "responsable", "nombre del actor"],
    "nombre_directivo": ["nombre directivo", "directivo", "dirigido a", "nombre del directivo", "nombre destinatario"],
    "prefijo": ["prefijo", "tratamiento", "titulo", "título"],
    "mesa": ["nombre de la mesa", "nombre mesa", "mesa", "tema", "asunto", "actividad"],
    "nivel": ["nivel", "compromiso", "tipo", "categoria", "categoría"],
    "fecha": ["fecha", "fecha mesa", "fecha programada", "dia", "día", "fecha de realizacion", "fecha de realización"],
    "dato": ["dato transformador", "dato", "transformador", "descripcion dato", "descripción dato", "resultado esperado"],
    "grupo": ["dependencia", "secretaria", "secretaría", "entidad", "despacho", "direccion", "dirección", "institucion", "institución", "grupo"],
    "firma_img": ["firma_img", "firma imagen", "firma", "imagen firma", "archivo firma", "firma path"],
    "logo_img":  ["logo_img", "logo imagen", "logo", "imagen logo", "archivo logo", "logo path"]
}

EXPECTED_HEADERS = ["nombre de la mesa", "nivel", "fecha", "dato transformador"]

def guess_mapping(df: pd.DataFrame) -> Dict[str, Optional[str]]:
    cols_norm = {c: _norm(c) for c in df.columns}
    mapping = {k: None for k in ["actor","nombre_directivo","prefijo","mesa","nivel","fecha","dato","grupo","firma_img","logo_img"]}
    for role, alias_list in SYNONYMS.items():
        exact = [c for c, cn in cols_norm.items() if cn in alias_list]
        if exact:
            mapping[role] = exact[0]; continue
        partial = [c for c, cn in cols_norm.items() if any(a in cn for a in alias_list)]
        if partial:
            mapping[role] = partial[0]
    return mapping

def parse_date(val) -> Optional[pd.Timestamp]:
    if pd.isna(val): return None
    if isinstance(val, (int, float)):
        try:
            return pd.to_datetime("1899-12-30") + pd.to_timedelta(int(val), unit="D")
        except Exception: pass
    for fmt in ("%Y-%m-%d","%d/%m/%Y","%d-%m-%Y","%Y/%m/%d","%m/%d/%Y","%d.%m.%Y"):
        try: return pd.to_datetime(datetime.strptime(str(val).strip(), fmt))
        except Exception: continue
    try: return pd.to_datetime(val, dayfirst=True, errors="coerce")
    except Exception: return None

def format_date_dmy(ts: Optional[pd.Timestamp]) -> str:
    if ts is None or pd.isna(ts): return ""
    try: return pd.to_datetime(ts).strftime("%d/%m/%Y")
    except Exception: return ""

def prepare_dataframe(df: pd.DataFrame, mapping: Dict[str, str]) -> pd.DataFrame:
    for key in ["actor","nombre_directivo","prefijo","mesa","nivel","fecha","dato"]:
        if key not in mapping or mapping[key] not in df.columns:
            raise ValueError(f"Columna requerida no encontrada para '{key}'.")
    cols = [
        mapping["actor"],
        mapping["nombre_directivo"],
        mapping["prefijo"],
        mapping.get("grupo"),
        mapping.get("firma_img"),
        mapping.get("logo_img"),
        mapping["mesa"], mapping["nivel"], mapping["fecha"], mapping["dato"]
    ]
    cols = [c for c in cols if c]
    work = df[cols].copy()
    out_cols = ["ACTOR","NOMBRE_DIRECTIVO","PREFIJO"] + (["GRUPO"] if mapping.get("grupo") else []) + \
               (["FIRMA_IMG"] if mapping.get("firma_img") else []) + (["LOGO_IMG"] if mapping.get("logo_img") else []) + \
               ["MESA","NIVEL","FECHA","DATO"]
    work.columns = out_cols
    work["_FECHA_TS"] = work["FECHA"].apply(parse_date)
    work["FECHA_FMT"] = work["_FECHA_TS"].apply(format_date_dmy)
    return work

def list_candidate_tables(doc: Document) -> List[int]:
    return [i for i, t in enumerate(doc.tables) if len(t.columns) == 4]

def _header_matches(table) -> bool:
    try: hdr_cells = table.rows[0].cells
    except Exception: return False
    headers = [_norm(c.text) for c in hdr_cells]
    return all(any(h == _norm(exp) or exp in h for h in headers) for exp in EXPECTED_HEADERS)

def find_target_table(doc: Document, prefer_index: int | None = None):
    if prefer_index is not None and 0 <= prefer_index < len(doc.tables):
        t = doc.tables[prefer_index]
        if len(t.columns) == 4: return t
    candidate = None
    for t in doc.tables:
        if _header_matches(t): return t
        if len(t.columns) == 4: candidate = t
    return candidate

def clear_table_keep_header(table)) -> None:
    while len(table.rows) > 1:
        tbl = table._tbl; tbl.remove(table.rows[-1]._tr)

def fill_table(table, rows: List[List[str]]) -> None:
    for r in rows:
        row = table.add_row()
        for i in range(min(4, len(r))):
            row.cells[i].text = "" if r[i] is None else str(r[i])

MESES_ES = ["enero","febrero","marzo","abril","mayo","junio","julio","agosto","septiembre","octubre","noviembre","diciembre"]
def month_name_es(month: int) -> str:
    if 1 <= month <= 12: return MESES_ES[month-1]
    return ""
