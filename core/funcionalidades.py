
# -*- coding: utf-8 -*-
from __future__ import annotations
import io, zipfile
from datetime import date
from typing import Dict, List, Optional, Tuple
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .backend import find_target_table, clear_table_keep_header, fill_table, slugify, month_name_es
from .routing import choose_template_for_group, render_derived_placeholders

def _expand_token_variants(key: str) -> List[str]:
    k1 = key
    k2 = key.replace(" ", "_")
    ups = list({k1.upper(), k2.upper()})
    return [f"{{{{{u}}}}}" for u in ups] + ups

def _replace_text_and_images(doc: Document, mapping_text: Dict[str, str], mapping_images: Dict[str, bytes], image_width_in: float = 1.5) -> None:
    tokens_text = {}
    for k, v in mapping_text.items():
        for var in _expand_token_variants(k):
            tokens_text[var] = str(v)

    tokens_img = {}
    for k, img_bytes in mapping_images.items():
        for var in _expand_token_variants(k):
            tokens_img[var] = img_bytes

    def _process_paragraph(p):
        full = "".join(run.text for run in p.runs)
        changed = False
        # imágenes
        for k, img_b in tokens_img.items():
            if k in full:
                full = full.replace(k, "")
                changed = True
                r = p.add_run()
                try:
                    from io import BytesIO
                    r.add_picture(BytesIO(img_b), width=Inches(image_width_in))
                except Exception:
                    pass
        # texto
        for k, v in tokens_text.items():
            if k in full:
                full = full.replace(k, v); changed = True
        if changed:
            for run in p.runs: run.text = ""
            p.add_run(full)

    for p in doc.paragraphs: _process_paragraph(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs: _process_paragraph(p)

def _rows_from_group(gdf: pd.DataFrame) -> List[List[str]]:
    return [[
        "" if pd.isna(r.MESA) else str(r.MESA),
        "" if pd.isna(r.NIVEL) else str(r.NIVEL),
        "" if pd.isna(r.FECHA_FMT) else str(r.FECHA_FMT),
        "" if pd.isna(r.DATO) else str(r.DATO),
    ] for r in gdf.itertuples(index=False)]

def _add_footer_with_pagenum(doc: Document, footer_text: str = "", logo_bytes: bytes | None = None, image_width_in: float = 1.0) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph("")
    if footer_text:
        p.add_run(footer_text + "  •  ")
    # "Página X de Y" con campos dinámicos
    run = p.add_run("Página ")
    _add_field(run, "PAGE")
    p.add_run(" de ")
    run2 = p.add_run("")
    _add_field(run2, "NUMPAGES")
    # Logo opcional al final
    if logo_bytes:
        from io import BytesIO
        p.add_run("   ")
        p.add_run().add_picture(BytesIO(logo_bytes), width=Inches(image_width_in))

def _add_field(run, field_code: str):
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), field_code)
    run._r.append(fld)

def generate_letters_per_group(
    work_df: pd.DataFrame,
    default_template_bytes: bytes,
    templates_map: Dict[str, bytes],
    routing_cfg: Dict,
    group_field: str = "ACTOR",
    table_index_default: int | None = None,
    newest_first: bool = True,
    city: str = "Medellín",
    letter_date: Optional[date] = None,
    naming_pattern: str = "CARTA_{GRUPO}.docx",
    image_assets: Dict[str, bytes] = None,
    image_width_in: float = 1.5,
) -> Tuple[Dict[str, bytes], Dict[str, str], pd.DataFrame]:
    outputs: Dict[str, bytes] = {}; errors: Dict[str, str] = {}; summary_rows: List[List[str]] = []
    work_df = work_df.sort_values("_FECHA_TS", ascending=not newest_first, na_position="last")
    d = letter_date or pd.Timestamp.today().date()
    fecha_larga = f"{city}, {d.day} de {month_name_es(d.month)} de {d.year}"

    derived_cfg = (routing_cfg or {}).get("derived_placeholders", {})

    for grp, gdf in work_df.groupby(group_field, dropna=False):
        grp_name = "(Sin grupo)" if pd.isna(grp) else str(grp)
        filas = _rows_from_group(gdf)
        try:
            # Regla y plantilla
            tpl_bytes, table_idx, rule = choose_template_for_group(grp_name, templates_map, routing_cfg)
            tpl_bytes = tpl_bytes or default_template_bytes
            table_idx = table_idx if table_idx is not None else table_index_default

            # Placeholders base
            actor = gdf["ACTOR"].dropna().astype(str).iloc[0] if "ACTOR" in gdf.columns and not gdf["ACTOR"].dropna().empty else grp_name
            nombre_dir = gdf["NOMBRE_DIRECTIVO"].dropna().astype(str).iloc[0] if "NOMBRE_DIRECTIVO" in gdf.columns and not gdf["NOMBRE_DIRECTIVO"].dropna().empty else ""
            prefijo = gdf["PREFIJO"].dropna().astype(str).iloc[0] if "PREFIJO" in gdf.columns and not gdf["PREFIJO"].dropna().empty else ""

            mapping_text = {"ACTOR": actor, "NOMBRE DIRECTIVO": nombre_dir, "PREFIJO": prefijo, "FECHA_CARTA": fecha_larga}
            # Derivados
            mapping_text.update(render_derived_placeholders(mapping_text, derived_cfg))

            # Imágenes por grupo
            img_map = {}
            if "FIRMA_IMG" in gdf.columns:
                fname = str(gdf["FIRMA_IMG"].dropna().astype(str).iloc[0]) if not gdf["FIRMA_IMG"].dropna().empty else None
                if fname and image_assets and fname in image_assets:
                    img_map["IMG_FIRMA"] = image_assets[fname]
            if "LOGO_IMG" in gdf.columns:
                fname = str(gdf["LOGO_IMG"].dropna().astype(str).iloc[0]) if not gdf["LOGO_IMG"].dropna().empty else None
                if fname and image_assets and fname in image_assets:
                    img_map["IMG_LOGO"] = image_assets[fname]

            # Construcción
            from io import BytesIO
            doc = Document(BytesIO(tpl_bytes))
            table = find_target_table(doc, prefer_index=table_idx)
            if table is None: raise RuntimeError("No se encontró una tabla válida (4 columnas) en la plantilla.")
            clear_table_keep_header(table); fill_table(table, filas)
            _replace_text_and_images(doc, mapping_text, img_map, image_width_in=image_width_in)

            # Footer auto
            footer_text = (routing_cfg or {}).get("footer_text", "")
            footer_logo_name = (routing_cfg or {}).get("footer_logo_name", None)
            footer_logo_bytes = image_assets.get(footer_logo_name) if (image_assets and footer_logo_name) else None
            if footer_text or footer_logo_bytes:
                _add_footer_with_pagenum(doc, footer_text=footer_text, logo_bytes=footer_logo_bytes, image_width_in=1.0)

            out = BytesIO(); doc.save(out); data = out.getvalue()
            safe_grp = slugify(grp_name)
            # Naming pattern por regla > global
            rule_namepat = rule.get("naming_pattern") if rule else None
            namepat = rule_namepat or naming_pattern
            fname = namepat.replace("{GRUPO}", safe_grp).replace("{ACTOR}", safe_grp)
            outputs[fname] = data; summary_rows.append([grp_name, len(filas)])
        except Exception as e:
            errors[grp_name] = str(e)

    index_df = pd.DataFrame(summary_rows, columns=["Grupo","Registros"]).sort_values("Grupo").reset_index(drop=True)
    return outputs, errors, index_df

def build_index_sheet(index_df: pd.DataFrame, errors: Dict[str, str]) -> bytes:
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as xlw:
        index_df.to_excel(xlw, sheet_name="Resumen", index=False)
        if errors:
            pd.DataFrame([{"Grupo":g,"Error":e} for g,e in errors.items()]).to_excel(xlw, sheet_name="Errores", index=False)
    return out.getvalue()

def make_zip(outputs: Dict[str, bytes]) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for fname, data in outputs.items():
            zf.writestr(fname, data)
    return buf.getvalue()
